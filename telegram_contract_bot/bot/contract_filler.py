from docx import Document
from docx.shared import Pt

def fill_contract(template_path: str, output_path: str, replacements: dict):
    doc = Document(template_path)

    def replace_preserving_format(paragraph, font_size=None):
        text = paragraph.text
        updated_text = text

        for key, val in replacements.items():
            updated_text = updated_text.replace(key, val)

        if updated_text != text:
            # Очищаем текущие runs
            for run in paragraph.runs:
                run.clear()

            # Добавляем новый текст с опциональным размером шрифта
            new_run = paragraph.add_run(updated_text)
            if font_size:
                new_run.font.size = Pt(font_size)

    # Заменяем в обычных абзацах (вне таблиц) без изменения размера
    for para in doc.paragraphs:
        replace_preserving_format(para)

    # Заменяем в таблицах, где отслеживается второй блок "Заказчик"
    customer_count = 0
    is_second_customer = False

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    # Отслеживаем блоки "Заказчик"
                    if "Заказчик" in para.text:
                        customer_count += 1
                        is_second_customer = (customer_count == 2)

                    # Применяем нужный размер
                    if is_second_customer:
                        replace_preserving_format(para, font_size=8)
                    else:
                        replace_preserving_format(para)

    # Сохраняем результат
    doc.save(output_path)